from fastapi import FastAPI, HTTPException, Depends, Header
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import List, Optionalfrom fastapi import FastAPI, HTTPException, Depends, Security
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, HttpUrl
from typing import List, Dict, Any, Optional
import httpx
import asyncio
import os
import json
import logging
from datetime import datetime
import hashlib
import numpy as np
from sentence_transformers import SentenceTransformer
import pinecone
from pinecone import Pinecone, ServerlessSpec
import openai
import fitz  # PyMuPDF
from docx import Document
import tempfile
import uvicorn
from contextlib import asynccontextmanager

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global variables for models and indices
embedding_model = None
pinecone_client = None
pinecone_index = None
document_chunks = []
document_metadata = []

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize models and resources on startup"""
    global embedding_model, pinecone_client, pinecone_index
    
    logger.info("Loading embedding model...")
    embedding_model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
    
    logger.info("Initializing Pinecone...")
    try:
        # Initialize Pinecone client
        pinecone_api_key = os.getenv("PINECONE_API_KEY")
        if not pinecone_api_key:
            logger.warning("Pinecone API key not found. Set PINECONE_API_KEY environment variable.")
        else:
            pinecone_client = Pinecone(api_key=pinecone_api_key)
            
            # Create or get index
            index_name = os.getenv("PINECONE_INDEX_NAME", "document-query-system")
            
            # Check if index exists, if not create it
            try:
                pinecone_index = pinecone_client.Index(index_name)
                logger.info(f"Connected to existing Pinecone index: {index_name}")
            except:
                logger.info(f"Creating new Pinecone index: {index_name}")
                pinecone_client.create_index(
                    name=index_name,
                    dimension=384,  # all-MiniLM-L6-v2 dimension
                    metric='cosine',
                    spec=ServerlessSpec(
                        cloud='aws',
                        region='us-east-1'
                    )
                )
                pinecone_index = pinecone_client.Index(index_name)
                
    except Exception as e:
        logger.error(f"Failed to initialize Pinecone: {e}")
        pinecone_client = None
        pinecone_index = None
    
    logger.info("Application startup complete")
    yield
    logger.info("Application shutdown")

app = FastAPI(
    title="LLM-Powered Document Query System",
    description="Intelligent query-retrieval system for insurance, legal, HR, and compliance documents",
    version="1.0.0",
    lifespan=lifespan
)

# Security
security = HTTPBearer()
BEARER_TOKEN = os.getenv("BEARER_TOKEN", "e0e272cd2f3ac51a8dda3c63908707c12fc63da7d51f0c7a8fbba91e80db3a88")

# OpenAI Configuration
openai.api_key = os.getenv("OPENAI_API_KEY")
if not openai.api_key:
    logger.warning("OpenAI API key not found. Set OPENAI_API_KEY environment variable.")

# Pinecone Configuration
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME", "document-query-system")

if not PINECONE_API_KEY:
    logger.warning("Pinecone API key not found. Set PINECONE_API_KEY environment variable.")

# Pydantic models
class DocumentQuery(BaseModel):
    documents: HttpUrl
    questions: List[str]

class QueryResponse(BaseModel):
    answers: List[str]

class ProcessingMetrics(BaseModel):
    processing_time: float
    token_usage: int
    chunks_processed: int
    similarity_scores: List[float]

# Authentication
async def verify_token(credentials: HTTPAuthorizationCredentials = Security(security)):
    if credentials.credentials != BEARER_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid authentication token")
    return credentials

class DocumentProcessor:
    """Handles document parsing and text extraction"""
    
    @staticmethod
    async def download_document(url: str) -> bytes:
        """Download document from URL"""
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url)
                response.raise_for_status()
                return response.content
        except Exception as e:
            logger.error(f"Error downloading document: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to download document: {str(e)}")

    @staticmethod
    def extract_text_from_pdf(content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                doc = fitz.open(temp_file.name)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                os.unlink(temp_file.name)
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise HTTPException(status_code=500, detail="Failed to extract text from PDF")

    @staticmethod
    def extract_text_from_docx(content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                doc = Document(temp_file.name)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                os.unlink(temp_file.name)
                return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise HTTPException(status_code=500, detail="Failed to extract text from DOCX")

    @classmethod
    async def process_document(cls, url: str) -> str:
        """Process document and extract text"""
        content = await cls.download_document(url)
        
        # Determine file type from URL or content
        url_lower = str(url).lower()
        if url_lower.endswith('.pdf') or b'%PDF' in content[:10]:
            return cls.extract_text_from_pdf(content)
        elif url_lower.endswith('.docx') or b'PK' in content[:2]:
            return cls.extract_text_from_docx(content)
        else:
            # Try to decode as text
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                raise HTTPException(status_code=400, detail="Unsupported file format")

class TextChunker:
    """Handles text chunking for embeddings"""
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks"""
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        current_size = 0
        
        for i, sentence in enumerate(sentences):
            sentence_size = len(sentence.split())
            
            if current_size + sentence_size > chunk_size and current_chunk:
                # Save current chunk
                chunks.append({
                    'text': current_chunk.strip(),
                    'start_sentence': max(0, i - len(current_chunk.split('. '))),
                    'end_sentence': i,
                    'word_count': current_size
                })
                
                # Start new chunk with overlap
                overlap_sentences = '. '.join(current_chunk.split('. ')[-overlap//10:])
                current_chunk = overlap_sentences + '. ' + sentence if overlap_sentences else sentence
                current_size = len(current_chunk.split())
            else:
                current_chunk += '. ' + sentence if current_chunk else sentence
                current_size += sentence_size
        
        # Add final chunk
        if current_chunk:
            chunks.append({
                'text': current_chunk.strip(),
                'start_sentence': max(0, len(sentences) - len(current_chunk.split('. '))),
                'end_sentence': len(sentences),
                'word_count': current_size
            })
        
        return chunks

class EmbeddingSearch:
    """Handles embedding generation and Pinecone operations"""
    
    @staticmethod
    def create_embeddings(texts: List[str]) -> List[List[float]]:
        """Create embeddings for text chunks"""
        global embedding_model
        if embedding_model is None:
            raise HTTPException(status_code=500, detail="Embedding model not initialized")
        
        embeddings = embedding_model.encode(texts, convert_to_tensor=False)
        return embeddings.tolist()
    
    @staticmethod
    def upsert_to_pinecone(embeddings: List[List[float]], chunks: List[Dict[str, Any]], doc_id: str) -> None:
        """Upsert embeddings and metadata to Pinecone"""
        global pinecone_index
        
        if pinecone_index is None:
            raise HTTPException(status_code=500, detail="Pinecone index not initialized")
        
        # Prepare vectors for upsert
        vectors = []
        for i, (embedding, chunk) in enumerate(zip(embeddings, chunks)):
            vector_id = f"{doc_id}_{i}"
            metadata = {
                'text': chunk['text'][:1000],  # Pinecone metadata size limit
                'start_sentence': chunk['start_sentence'],
                'end_sentence': chunk['end_sentence'],
                'word_count': chunk['word_count'],
                'doc_id': doc_id,
                'chunk_index': i
            }
            vectors.append({
                'id': vector_id,
                'values': embedding,
                'metadata': metadata
            })
        
        # Upsert in batches to handle large documents
        batch_size = 100
        for i in range(0, len(vectors), batch_size):
            batch = vectors[i:i + batch_size]
            pinecone_index.upsert(vectors=batch)
    
    @staticmethod
    def search_similar(query: str, doc_id: str, k: int = 5) -> List[Dict[str, Any]]:
        """Search for similar text chunks in Pinecone"""
        global embedding_model, pinecone_index
        
        if embedding_model is None or pinecone_index is None:
            raise HTTPException(status_code=500, detail="Search components not initialized")
        
        # Create query embedding
        query_embedding = embedding_model.encode([query], convert_to_tensor=False)
        query_embedding = query_embedding.tolist()[0]
        
        # Search in Pinecone with filter for the specific document
        search_results = pinecone_index.query(
            vector=query_embedding,
            top_k=k,
            include_metadata=True,
            filter={"doc_id": doc_id}
        )
        
        results = []
        for match in search_results['matches']:
            chunk_data = {
                'text': match['metadata']['text'],
                'start_sentence': match['metadata']['start_sentence'],
                'end_sentence': match['metadata']['end_sentence'],
                'word_count': match['metadata']['word_count']
            }
            
            results.append({
                'chunk': chunk_data,
                'similarity_score': float(match['score']),
                'id': match['id']
            })
        
        return results

class LLMProcessor:
    """Handles LLM queries and response generation"""
    
    @staticmethod
    async def generate_answer(question: str, context_chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate answer using OpenAI GPT"""
        if not openai.api_key:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        # Prepare context
        context = "\n\n".join([
            f"Context {i+1} (Score: {chunk['similarity_score']:.3f}):\n{chunk['chunk']['text']}"
            for i, chunk in enumerate(context_chunks[:3])  # Use top 3 chunks
        ])
        
        # Create prompt
        prompt = f"""Based on the following context from insurance policy documents, please answer the question accurately and concisely.

Context:
{context}

Question: {question}

Instructions:
1. Answer based only on the provided context
2. If the information is not available in the context, state "Information not available in the provided context"
3. Provide specific details like waiting periods, coverage limits, conditions, etc. when available
4. Be precise and factual

Answer:"""

        try:
            response = await openai.ChatCompletion.acreate(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert insurance policy analyzer. Provide accurate, concise answers based on policy documents."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=300,
                temperature=0.1
            )
            
            answer = response.choices[0].message.content.strip()
            token_usage = response.usage.total_tokens
            
            return {
                'answer': answer,
                'token_usage': token_usage,
                'context_used': len(context_chunks),
                'similarity_scores': [chunk['similarity_score'] for chunk in context_chunks]
            }
            
        except Exception as e:
            logger.error(f"Error generating answer: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to generate answer: {str(e)}")

# API Routes
@app.get("/")
async def root():
    return {"message": "LLM-Powered Document Query System", "status": "running"}

@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}

@app.post("/hackrx/run", response_model=QueryResponse)
async def run_query_system(
    request: DocumentQuery,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
) -> QueryResponse:
    """Main endpoint for document querying"""
    global document_chunks, document_metadata
    
    start_time = datetime.utcnow()
    
    try:
        logger.info(f"Processing document: {request.documents}")
        
        # Generate unique document ID
        doc_id = hashlib.md5(str(request.documents).encode()).hexdigest()[:8]
        
        # Step 1: Download and extract text
        document_text = await DocumentProcessor.process_document(str(request.documents))
        logger.info(f"Extracted {len(document_text)} characters from document")
        
        # Step 2: Chunk text
        document_chunks = TextChunker.chunk_text(document_text)
        logger.info(f"Created {len(document_chunks)} text chunks")
        
        # Step 3: Create embeddings and upsert to Pinecone
        chunk_texts = [chunk['text'] for chunk in document_chunks]
        embeddings = EmbeddingSearch.create_embeddings(chunk_texts)
        EmbeddingSearch.upsert_to_pinecone(embeddings, document_chunks, doc_id)
        logger.info("Upserted embeddings to Pinecone")
        
        # Step 4: Process questions
        answers = []
        total_tokens = 0
        
        for question in request.questions:
            logger.info(f"Processing question: {question[:50]}...")
            
            # Search for relevant chunks in Pinecone
            similar_chunks = EmbeddingSearch.search_similar(question, doc_id, k=5)
            
            # Generate answer
            result = await LLMProcessor.generate_answer(question, similar_chunks)
            answers.append(result['answer'])
            total_tokens += result['token_usage']
        
        processing_time = (datetime.utcnow() - start_time).total_seconds()
        logger.info(f"Completed processing in {processing_time:.2f}s, used {total_tokens} tokens")
        
        return QueryResponse(answers=answers)
        
    except Exception as e:
        logger.error(f"Error processing request: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/hackrx/metrics")
async def get_processing_metrics(
    request: DocumentQuery,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
) -> ProcessingMetrics:
    """Get detailed processing metrics"""
    start_time = datetime.utcnow()
    
    # Simulate processing (in real implementation, this would be actual processing)
    await asyncio.sleep(0.1)  # Simulate processing time
    
    processing_time = (datetime.utcnow() - start_time).total_seconds()
    
    return ProcessingMetrics(
        processing_time=processing_time,
        token_usage=150,  # Placeholder
        chunks_processed=len(document_chunks) if document_chunks else 0,
        similarity_scores=[0.85, 0.72, 0.65, 0.58, 0.45]  # Placeholder
    )

# For Gunicorn deployment, we don't need the uvicorn.run() call
# The app instance will be imported directly by Gunicorn

if __name__ == "__main__":
    # This is only for local development
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=port,
        reload=False,
        log_level="info"
    )
import requests
import os
from datetime import datetime
import logging
from sentence_transformers import SentenceTransformer
import pinecone
import PyPDF2
import docx
import re
import json
import openai
from nltk.tokenize import sent_tokenize

# Initialize FastAPI app
app = FastAPI(
    title="Intelligent Query-Retrieval System",
    description="LLM-powered system for processing documents and answering queries",
    version="1.0.0"
)

# Security
security = HTTPBearer()
TEAM_TOKEN = "e0e272cd2f3ac51a8dda3c63908707c12fc63da7d51f0c7a8fbba91e80db3a88"

# Models
class DocumentInput(BaseModel):
    documents: str
    questions: List[str]

class QueryResponse(BaseModel):
    answers: List[str]

# Configuration
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
CHUNK_SIZE = 500  # characters
OVERLAP = 50      # characters
MAX_TOKENS = 4096  # For GPT-4
TEMPERATURE = 0.3
PINECONE_INDEX_NAME = "intelligent-query-system"
EMBEDDING_DIMENSION = 384  # For all-MiniLM-L6-v2

# Initialize components
embedding_model = SentenceTransformer(EMBEDDING_MODEL)
pinecone.init(api_key=os.getenv("PINECONE_API_KEY"), environment=os.getenv("PINECONE_ENV"))

# Check if index exists, create if not
if PINECONE_INDEX_NAME not in pinecone.list_indexes():
    pinecone.create_index(
        name=PINECONE_INDEX_NAME,
        dimension=EMBEDDING_DIMENSION,
        metric="cosine"
    )

# Connect to index
index = pinecone.Index(PINECONE_INDEX_NAME)

# Utility Functions
def download_file(url: str) -> str:
    """Download document from URL"""
    local_path = f"temp_{datetime.now().timestamp()}"
    if url.endswith('.pdf'):
        local_path += '.pdf'
    elif url.endswith('.docx'):
        local_path += '.docx'
    else:
        raise ValueError("Unsupported file format")
    
    with requests.get(url, stream=True) as r:
        r.raise_for_status()
        with open(local_path, 'wb') as f:
            for chunk in r.iter_content(chunk_size=8192):
                f.write(chunk)
    return local_path

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF"""
    text = ""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX"""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def preprocess_text(text: str) -> str:
    """Clean and normalize text"""
    text = re.sub(r'\s+', ' ', text)  # Replace multiple whitespace
    text = text.strip()
    return text

def chunk_text(text: str) -> List[str]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + CHUNK_SIZE
        chunks.append(text[start:end])
        start = end - OVERLAP
    return chunks

def upsert_to_pinecone(chunks: List[str], document_id: str):
    """Upload document chunks to Pinecone with embeddings"""
    vectors = []
    for i, chunk in enumerate(chunks):
        # Generate embedding for each chunk
        embedding = embedding_model.encode(chunk).tolist()
        
        # Create unique ID for each chunk
        chunk_id = f"{document_id}-{i}"
        
        vectors.append((chunk_id, embedding, {"text": chunk, "document_id": document_id}))
    
    # Upsert in batches of 100 (Pinecone limit)
    for i in range(0, len(vectors), 100):
        batch = vectors[i:i+100]
        index.upsert(vectors=batch)

def semantic_search(query: str, k: int = 3) -> List[dict]:
    """Perform semantic search using Pinecone"""
    query_embedding = embedding_model.encode(query).tolist()
    
    results = index.query(
        vector=query_embedding,
        top_k=k,
        include_metadata=True
    )
    
    return [
        {
            "chunk": match["metadata"]["text"],
            "score": match["score"],
            "metadata": match["metadata"]
        }
        for match in results["matches"]
    ]

def generate_response_with_llm(query: str, context: str) -> str:
    """Generate response using LLM with context"""
    prompt = f"""
    You are an intelligent query answering system for insurance, legal, HR, and compliance documents.
    Based on the following document context, answer the user's question precisely.
    If the information is not in the document, say "The document does not specify."

    Document Context:
    {context}

    Question: {query}
    Answer:"""
    
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": prompt}],
        temperature=TEMPERATURE,
        max_tokens=MAX_TOKENS // 2
    )
    
    return response.choices[0].message.content.strip()

def process_document(document_url: str):
    """Process document and upload to Pinecone"""
    try:
        local_path = download_file(document_url)
        
        if document_url.endswith('.pdf'):
            text = extract_text_from_pdf(local_path)
        elif document_url.endswith('.docx'):
            text = extract_text_from_docx(local_path)
        else:
            raise ValueError("Unsupported file format")
        
        text = preprocess_text(text)
        chunks = chunk_text(text)
        
        # Use document URL as document ID
        document_id = hashlib.md5(document_url.encode()).hexdigest()
        upsert_to_pinecone(chunks, document_id)
        
        os.remove(local_path)
        return True
    except Exception as e:
        logging.error(f"Error processing document: {str(e)}")
        return False

# API Endpoints
@app.post("/api/v1/hackrx/run", response_model=QueryResponse)
async def run_query(
    input_data: DocumentInput,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Process document and answer questions"""
    # Authentication
    if credentials.credentials != TEAM_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    # Process document
    if not process_document(input_data.documents):
        raise HTTPException(status_code=400, detail="Document processing failed")
    
    # Answer questions
    answers = []
    for question in input_data.questions:
        try:
            # Semantic search for relevant chunks
            search_results = semantic_search(question)
            context = "\n\n".join([res['chunk'] for res in search_results[:3]])
            
            # Generate answer with LLM
            answer = generate_response_with_llm(question, context)
            answers.append(answer)
        except Exception as e:
            answers.append(f"Error processing question: {str(e)}")
    
    return {"answers": answers}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

